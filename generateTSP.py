#!/bin/python2.7

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def textToCell( doc, targetCell, string, align, bold ):
	style = doc.styles['Normal']
	font = style.font
	font.name = 'Calibri'
	font.size = Pt(11)

	targetCell.text = ''
	p = targetCell.paragraphs[0]
	p.style = doc.styles['Normal']
	p.alignment = align
        r = p.add_run(string)
	r.bold = bold

def generateTSP( module_name, version_number ):
	tspDoc = Document()
	table = tspDoc.add_table(rows=0, cols=2)
	table.style = 'Table Grid'
	# Add Title
	titleRow = table.add_row()
	titleCell = table.cell(0,0)
	titleCell.merge(table.cell(0,1))
	textToCell(tspDoc, titleCell,'PROJECT INFORMATION', WD_ALIGN_PARAGRAPH.CENTER, True)
	# Add Remaining Rows
	for x in range(4):
		table.add_row()
	# Add Project ID
	textToCell(tspDoc, table.cell(1,0), 'Project ID', WD_ALIGN_PARAGRAPH.LEFT, False)
        table.cell(1,1).text = 'RIU100:' + module_name + '_TSP.CA-TSP_SCR'
	# Add Test Specification Ref
	textToCell(tspDoc, table.cell(2,0), 'Test Specification Ref.', WD_ALIGN_PARAGRAPH.LEFT, False)
	table.cell(2,1).text = 'RIU100:' + module_name + '_TSP.CA-TSP_SCR'
	# Add Test Specification Issue
	textToCell(tspDoc, table.cell(3,0), 'Test Specification Issue', WD_ALIGN_PARAGRAPH.LEFT, False)
        table.cell(3,1).text = 'ce#' + str(version_number)
	# Add Test Specification Revision
	textToCell(tspDoc, table.cell(4,0), 'Test Specification Revision', WD_ALIGN_PARAGRAPH.LEFT, False)

	tspDoc.save(module_name + 'TSP_REVIEW.docx')


generateTSP('test', 1);
